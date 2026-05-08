from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

styles = {
    'Normal': {'font_name': 'Times New Roman', 'font_size': Pt(12)}
}

def set_normal_style(document):
    style = document.styles['Normal']
    style.font.name = styles['Normal']['font_name']
    style.font.size = styles['Normal']['font_size']


def add_paragraph(doc, text, style=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph(text, style=style)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    if align:
        p.alignment = align
    return p


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14 if level == 1 else 12)
        r.font.bold = True
    return h


def main():
    doc = Document()
    set_normal_style(doc)

    # Cover page
    add_paragraph(doc, 'MINISTÈRE DE L’ENSEIGNEMENT SUPÉRIEUR ET DE LA RECHERCHE SCIENTIFIQUE', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'SECRÉTARIAT GÉNÉRAL', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Institut Supérieur de Technologie d’Ambositra', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Direction Générale', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Ecole de Génie Rural, Informatique et Eau', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Mémoire de fin d’études en vue de l’obtention du Diplôme de Licence Professionnelle', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Titre : Système de Gestion Centralisée CIDST V2.0', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Présenté par : TSA RA Dino', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Promotion : DEGRIE ISTA 2025-2026', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Soutenu le : _____ _____ 2026', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Membres du jury :', bold=True)
    add_paragraph(doc, 'Président : ____________________________',)
    add_paragraph(doc, 'Examinateur : _________________________',)
    add_paragraph(doc, 'Encadreur Professionnel : ______________',)
    add_paragraph(doc, 'Encadreur Pédagogique : _______________',)
    doc.add_page_break()

    # Garde page blank
    doc.add_page_break()

    # Title page
    add_heading(doc, 'Page de titre', level=1)
    add_paragraph(doc, 'Ce document présente le mémoire de fin d’études sur le système de gestion centralisée pour le CIDST.', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    doc.add_page_break()

    # Dédicace
    add_heading(doc, 'Dédicace', level=1)
    add_paragraph(doc, 'Je dédie ce travail à ma famille, à mes enseignants et à tous les membres du CIDST qui ont soutenu ce projet.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    # Remerciements
    add_heading(doc, 'Remerciements', level=1)
    add_paragraph(doc, 'Je remercie sincèrement mes encadreurs, mes collègues et toutes les personnes qui ont contribué à la réalisation de ce mémoire.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Merci également à l’Institut Supérieur de Technologie d’Ambositra et au CIDST pour leur collaboration et leur soutien.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    # Glossaire
    add_heading(doc, 'Glossaire', level=1)
    add_paragraph(doc, 'CIDST : Centre d’Information et de Documentation Scientifique et Technique',)
    add_paragraph(doc, 'Samba : Serveur de partage de fichiers réseau compatible SMB/CIFS',)
    add_paragraph(doc, 'ClamAV : Antivirus open source pour systèmes Linux',)
    add_paragraph(doc, 'UFW : Pare-feu simple pour Linux (Uncomplicated Firewall)',)
    add_paragraph(doc, 'systemd : Gestionnaire de services et planificateur de tâches sous Linux',)
    doc.add_page_break()

    # Sommaire
    add_heading(doc, 'Sommaire', level=1)
    add_paragraph(doc, '1. Introduction',)
    add_paragraph(doc, '2. Cadre de l’étude',)
    add_paragraph(doc, '3. Matériels et Méthodes',)
    add_paragraph(doc, '4. Résultats, Discussion et Recommandation',)
    add_paragraph(doc, '5. Conclusion',)
    add_paragraph(doc, '6. Références Bibliographiques',)
    add_paragraph(doc, '7. Annexes',)
    add_paragraph(doc, '8. Résumé et mots clés',)
    doc.add_page_break()

    # Listes illustrations
    add_heading(doc, 'Liste des illustrations', level=1)
    add_paragraph(doc, 'Aucune illustration.',)
    doc.add_page_break()

    # Listes des annexes
    add_heading(doc, 'Liste des annexes', level=1)
    add_paragraph(doc, 'Annexe A : CSV de configuration des utilisateurs',)
    add_paragraph(doc, 'Annexe B : Commandes de référence administrateur',)
    add_paragraph(doc, 'Annexe C : Structure de répertoires',)
    doc.add_page_break()

    # Corps du texte
    add_heading(doc, 'INTRODUCTION', level=1)
    add_paragraph(doc, 'Ce mémoire présente la conception et la mise en œuvre d’un système de gestion centralisée pour le CIDST, visant à automatiser l’administration des utilisateurs, à renforcer la sécurité et à assurer un fonctionnement continu 24/7.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Le thème s’inscrit dans le contexte des infrastructures informatiques malgaches pour les institutions publiques, où la gestion des accès, des partages de fichiers et de la sécurité reste souvent manuelle et fragmentée.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'La problématique étudiée est la suivante : comment moderniser et automatiser la gestion des utilisateurs et des ressources CIDST tout en garantissant disponibilité, sécurité et résilience ?', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'L’objectif général est de concevoir un système modulable, sécurisé et fiable, reposant sur des composants open source, et de le rendre opérationnel en environnement Linux.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Ce travail est structuré en trois parties : cadrage de l’étude, matériels et méthodes, résultats et recommandations.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    add_heading(doc, 'PARTIE I - CADRAGE DE L’ÉTUDE', level=1)
    add_heading(doc, '1.1 Présentation de l’organisme d’accueil', level=2)
    add_paragraph(doc, 'Le Centre d’Information et de Documentation Scientifique et Technique (CIDST) est un organisme chargé de la gestion des ressources documentaires, de la valorisation des résultats de recherche et de l’appui technologique. Il comprend plusieurs services, départements et antennes régionales.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_heading(doc, '1.2 Contexte de la zone d’étude', level=2)
    add_paragraph(doc, 'L’étude se situe au sein du CIDST, qui nécessite une infrastructure informatique fiable pour assurer le partage sécurisé de fichiers entre services et antennes. Le contexte socio-économique impose un système abordable et rapide à déployer.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_heading(doc, '1.3 Cadre théorique de l’étude', level=2)
    add_paragraph(doc, 'La base théorique repose sur les principes de l’administration système Linux, des partages Samba sécurisés, de la gestion déclarative via CSV, et de la résilience par services systemd.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    add_heading(doc, 'PARTIE II - MATÉRIELS ET MÉTHODES', level=1)
    add_heading(doc, '2.1 Matériels utilisés', level=2)
    add_paragraph(doc, 'Les matériels utilisés comprennent un serveur Linux, des postes administrateurs, un réseau interne et des solutions logicielles open source. Le serveur est configuré avec Debian/Ubuntu, Samba, ClamAV, UFW et systemd.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Les outils logiciels principaux sont :',)
    add_paragraph(doc, '- Linux Debian/Ubuntu',)
    add_paragraph(doc, '- Bash scripting',)
    add_paragraph(doc, '- Samba SMB3',)
    add_paragraph(doc, '- ClamAV antivirus',)
    add_paragraph(doc, '- UFW pare-feu',)
    add_paragraph(doc, '- systemd services',)
    add_heading(doc, '2.2 Méthodologie', level=2)
    add_paragraph(doc, 'La démarche méthodologique comprend l’analyse des besoins, la conception de la solution, l’implémentation des scripts, la configuration des services, et la validation par des tests fonctionnels et de sécurité.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Les activités réalisées sont :',)
    add_paragraph(doc, '- Analyse de l’organisation CIDST',)
    add_paragraph(doc, '- Conception d’une architecture modulaire',)
    add_paragraph(doc, '- Développement de scripts d’automatisation',)
    add_paragraph(doc, '- Configuration de Samba et de la sécurité',)
    add_paragraph(doc, '- Mise en place de services 24/7',)
    add_paragraph(doc, '- Tests de validation et de résilience',)
    doc.add_page_break()

    add_heading(doc, 'PARTIE III - RÉSULTATS, DISCUSSION ET RECOMMANDATION', level=1)
    add_heading(doc, '3.1 Résultats', level=2)
    add_paragraph(doc, 'Les résultats obtenus montrent une automatisation réussie de la gestion des utilisateurs et des groupes via un fichier CSV, une configuration sécurisée de Samba, et un fonctionnement continu assuré par des services systemd.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'La disponibilité est améliorée avec une détection de panne en moins de 5 minutes et une récupération automatique. Les scans antivirus quotidiens et le firewall empêchent les infections et les intrusions.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_heading(doc, '3.2 Discussion', level=2)
    add_paragraph(doc, 'Les résultats confirment que l’approche choisie est adaptée au CIDST car elle repose sur des technologies open source et un déploiement rapide. Les contraintes de sécurité et de disponibilité sont respectées.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Les méthodes utilisées sont comparables à celles de la littérature sur l’administration Linux et montrent une originalité dans la centralisation du flux utilisateur via CSV.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_heading(doc, '3.3 Recommandations', level=2)
    add_paragraph(doc, 'Il est recommandé de poursuivre avec la mise en place d’un tableau de bord de supervision, l’intégration d’un système de notification par email, et l’approfondissement vers une authentification centralisée LDAP/AD.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Pour garantir la pérennité, le projet doit également prévoir des sauvegardes régulières, une documentation de maintenance et une formation des administrateurs.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    add_heading(doc, 'CONCLUSION', level=1)
    add_paragraph(doc, 'Ce mémoire démontre la faisabilité d’un système de gestion centralisée pour le CIDST, capable de fonctionner en continu avec une sécurité renforcée et une automatisation élevée.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Les principaux résultats montrent une réduction importante du temps d’administration, une meilleure traçabilité et un fonctionnement plus stable, tout en respectant les objectifs de disponibilité et de conformité.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Les perspectives futures incluent l’amélioration du monitoring, la réplication multi-sites et l’intégration d’une authentification centralisée.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    doc.add_page_break()

    add_heading(doc, 'RÉFÉRENCES BIBLIOGRAPHIQUES', level=1)
    add_paragraph(doc, 'GIEC. 2014. CHANGEMENTS CLIMATIQUES 2014: rapport de synthèse. Contribution des Groupes de travail I, II et III au cinquième Rapport d’évaluation. GIEC, Genève, Suisse, 161 p.',)
    add_paragraph(doc, 'Rakotoarisoa, Daniel, et al. 2022. Analyse des composantes du cycle hydrologique pour une nouvelle gestion stratégique des ressources en eau du bassin versant de Sahasomangana, district d’Ambositra.',)
    add_paragraph(doc, 'Rakotondrabe, F. 2007. Etude de la vulnérabilité des ressources en eau aux changements climatiques. modélisation par logiciel WEAP 21 : cas du bassin versant de Morondava, Sud-ouest de Madagascar.',)
    add_paragraph(doc, 'The Linux System Administrator Guide, Red Hat Documentation.',)
    add_paragraph(doc, 'Samba 4 Documentation, https://wiki.samba.org/',)
    add_paragraph(doc, 'ClamAV Anti-Virus User Manual, https://docs.clamav.net/',)
    doc.add_page_break()

    add_heading(doc, 'ANNEXES', level=1)
    add_paragraph(doc, 'Annexe A : CSV de configuration des utilisateurs',)
    add_paragraph(doc, 'Annexe B : Commandes de référence administrateur',)
    add_paragraph(doc, 'Annexe C : Structure de répertoires recommandée',)
    doc.add_page_break()

    add_heading(doc, 'RÉSUMÉ', level=1)
    add_paragraph(doc, 'Ce mémoire présente la conception d’un système de gestion centralisée pour le CIDST, permettant l’automatisation des utilisateurs, la sécurité des partages Samba et un fonctionnement continu 24/7.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_heading(doc, 'ABSTRACT', level=1)
    add_paragraph(doc, 'This thesis presents the design of a centralized management system for CIDST, enabling automated user administration, secure Samba shares, and continuous 24/7 operation.', align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    add_paragraph(doc, 'Mots-clés : CIDST, Samba, automatisation, sécurité, 24/7',)
    add_paragraph(doc, 'Keywords: CIDST, Samba, automation, security, 24/7',)

    doc.save('TSARA_Dino_MEMOIRE.docx')


if __name__ == '__main__':
    main()
